from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import datetime

doc = Document()

# ── Estilos ──────────────────────────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Calibri"
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── Portada ──────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "Documentación Técnica\nSistema de Análisis de Encuestas de Satisfacción"
)
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(
    "Aplicación Web para la Evaluación del Desempeño\nDocente, Facilitador, Supervisor y Asignatura"
)
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

for _ in range(4):
    doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run(f"Elaborado: {datetime.date.today().strftime('%d/%m/%Y')}")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_page_break()

# ── Tabla de contenidos (manual) ────────────────────────────────────────
doc.add_heading("Índice", level=1)
toc_items = [
    "1. Resumen Ejecutivo",
    "2. Descripción General",
    "3. Tecnologías Utilizadas",
    "4. Estructura del Proyecto",
    "5. Flujo de la Aplicación",
    "6. Formato Esperado de Archivos Excel",
    "7. Modelo de Puntuación",
    "8. Dimensiones Evaluadas por Rol",
    "9. Páginas y Funcionalidades",
    "10. Limitaciones y Consideraciones",
    "11. Recomendaciones",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# 1. RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("1. Resumen Ejecutivo", level=1)
doc.add_paragraph(
    "El presente sistema es una aplicación web desarrollada en Python utilizando Streamlit "
    "que permite analizar los resultados de encuestas de satisfacción estudiantil exportadas "
    "desde la plataforma Moodle en formato Microsoft Excel (.xlsx). La aplicación procesa "
    "evaluaciones correspondientes a cuatro roles académicos: Docente, Facilitador, Supervisor "
    "y Asignatura, cada uno evaluado por distintos actores (estudiantes, equipos de carrera, "
    "supervisores, facilitadores y equipo pedagógico)."
)
doc.add_paragraph(
    "El sistema ofrece visualizaciones interactivas, tablas comparativas, gráficos de radar, "
    "barras apiladas y mapas de calor, permitiendo a los usuarios identificar fortalezas y "
    "debilidades en distintas dimensiones pedagógicas. Todo el procesamiento se realiza en "
    "memoria durante la sesión del usuario, sin necesidad de bases de datos ni conexión a "
    "internet (una vez cargada la página)."
)

# ══════════════════════════════════════════════════════════════════════════
# 2. DESCRIPCIÓN GENERAL
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("2. Descripción General", level=1)
doc.add_paragraph(
    "La aplicación permite a los usuarios cargar archivos Excel provenientes de las encuestas "
    "de satisfacción de Moodle, organizarlos según el tipo de evaluado (Docente, Facilitador, "
    "Supervisor o Asignatura) y el evaluador correspondiente, y luego visualizar los resultados "
    "a través de las siguientes secciones:"
)
pages = [
    (
        "Cargar Datos",
        "Interfaz para subir archivos Excel organizados por tipo de evaluado y evaluador.",
    ),
    (
        "Resultados",
        "Vista detallada de una evaluación individual con métricas, gráficos y comentarios.",
    ),
    (
        "Comparación",
        "Agrupa todas las evaluaciones cargadas por módulo, promedia puntuaciones y muestra rankings.",
    ),
    (
        "Diagnóstico General",
        "Panel consolidado con métricas globales, ranking completo, análisis por dimensión y tipo, y tabla de fortalezas/debilidades.",
    ),
]
for name, desc in pages:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "Cada evaluación se compone de un conjunto de preguntas agrupadas en dimensiones pedagógicas. "
    "Las respuestas se califican en una escala de 3 puntos (Sí, A veces, No), y el sistema calcula "
    "un puntaje promedio ponderado para cada pregunta, cada dimensión y el total general."
)

# ══════════════════════════════════════════════════════════════════════════
# 3. TECNOLOGÍAS UTILIZADAS
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("3. Tecnologías Utilizadas", level=1)

techs = [
    ("Python 3.14", "Lenguaje de programación principal."),
    (
        "Streamlit >= 1.35.0",
        "Framework web para crear la interfaz de usuario interactiva. Permite crear dashboards con solo Python, manejando el ciclo request-response de forma transparente.",
    ),
    (
        "Pandas >= 2.0.0",
        "Biblioteca de manipulación de datos. Se utiliza para estructurar los datos extraídos, crear DataFrames, generar resúmenes estadísticos y exportar a CSV.",
    ),
    (
        "Plotly >= 5.18.0",
        "Biblioteca de visualización interactiva. Genera gráficos de barras, barras apiladas al 100%, gráficos circulares (pie charts), gráficos de radar, mapas de calor y más.",
    ),
    (
        "OpenPyXL >= 3.1.0",
        "Biblioteca para lectura y escritura de archivos Excel .xlsx. Se utiliza para leer los archivos exportados desde Moodle en modo read-only para eficiencia.",
    ),
    (
        "python-docx",
        "Biblioteca utilizada para generar la presente documentación en formato Word.",
    ),
]
table = doc.add_table(rows=1, cols=2)
table.style = "Light Shading Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Tecnología"
hdr[1].text = "Descripción"
for name, desc in techs:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = desc

doc.add_paragraph()

dep_titles = [
    (
        "io.BytesIO",
        "Módulo estándar de Python. Se utiliza para tratar los archivos subidos como objetos binarios en memoria.",
    ),
    (
        "re (expresiones regulares)",
        "Para extraer números de preguntas y procesar metadatos.",
    ),
    (
        "unicodedata",
        "Para normalizar caracteres Unicode (eliminación de tildes y normalización NFKD).",
    ),
    (
        "collections.defaultdict",
        "Para agrupar datos por módulo sin necesidad de verificar existencia de claves.",
    ),
]
doc.add_paragraph(
    "Dependencias de la biblioteca estándar de Python:", style="List Bullet"
)
for name, desc in dep_titles:
    p = doc.add_paragraph()
    run = p.add_run(f"{name}: ")
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════════════════
# 4. ESTRUCTURA DEL PROYECTO
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("4. Estructura del Proyecto", level=1)
doc.add_paragraph(
    "El proyecto está contenido en un único archivo principal de aplicación más archivos auxiliares:"
)
code_block = (
    "encuestas_app/\n"
    "├── app_completa.py          # Aplicación principal (~1550 líneas)\n"
    "├── requirements.txt          # Dependencias del proyecto\n"
    "├── README.md                 # Instrucciones de uso\n"
    "├── generar_documentacion.py  # Script para generar este documento\n"
    "├── Encuestas de Satisfacción/ # Carpeta con archivos Excel de ejemplo\n"
    "│   ├── Asignatura ev. por el Docente/\n"
    "│   ├── Docente ev. por el Equipo de la Carrera/\n"
    "│   ├── Docente ev. por el Estudiante/\n"
    "│   ├── Facilitador ev. por el Supervisor/\n"
    "│   ├── Facilitador ev. por los Estudiantes/\n"
    "│   ├── Supervisor ev. por el equipo Pedagógico/\n"
    "│   └── Supervisor ev. por el Facilitador/\n"
    "└── venv/                     # Entorno virtual de Python\n"
)
p = doc.add_paragraph()
run = p.add_run(code_block)
run.font.name = "Consolas"
run.font.size = Pt(9)

doc.add_paragraph(
    "El archivo app_completa.py contiene toda la lógica de la aplicación: desde la definición "
    "de tipos de evaluación y dimensiones, pasando por el parseo de archivos Excel, hasta "
    "el renderizado de las páginas y gráficos."
)

# ══════════════════════════════════════════════════════════════════════════
# 5. FLUJO DE LA APLICACIÓN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("5. Flujo de la Aplicación", level=1)

doc.add_heading("5.1 Inicio", level=2)
doc.add_paragraph(
    "Al ejecutar streamlit run app_completa.py, la aplicación inicia y muestra una interfaz "
    "con un menú de navegación en la parte superior que permite alternar entre las cuatro "
    "secciones principales."
)

doc.add_heading("5.2 Carga de Datos", level=2)
doc.add_paragraph("El usuario debe seguir estos pasos para cargar datos:")
steps = [
    "Seleccionar el tipo de evaluado (Docente, Facilitador, Supervisor o Asignatura).",
    "Seleccionar el evaluador correspondiente (las opciones cambian según el tipo seleccionado).",
    "Utilizar el cargador de archivos (st.file_uploader) para subir uno o más archivos Excel.",
    "El sistema procesa cada archivo mediante la función parse_evaluacion(), que extrae metadatos, preguntas y comentarios.",
    "Los resultados se almacenan en st.session_state.data y están disponibles para todas las páginas.",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("5.3 Procesamiento de Archivos", level=2)
doc.add_paragraph(
    "La función parse_evaluacion() es el núcleo del sistema. Realiza las siguientes operaciones:"
)
proc_steps = [
    "Abre el archivo Excel con openpyxl en modo read_only para eficiencia.",
    "Itera sobre todas las filas de la hoja activa.",
    "Extrae metadatos (carrera, sede, módulo, grupo, docente, fechas, respuestas) de las primeras filas.",
    'Localiza la fila que contiene "Etiqueta" para identificar el inicio de los datos de preguntas.',
    "Para cada fila de pregunta, extrae los valores de las columnas Excelentes (Sí), Muy Buenos (A veces) y Bueno (No).",
    "Agrupa preguntas duplicadas (solo para Docente evaluado por Estudiante) promediando sus valores.",
    'Detecta el inicio de la sección de comentarios cuando el texto de la pregunta contiene "observaciones" o "comentarios".',
    "Filtra comentarios irrelevantes mediante una lista de palabras vacías (stop-words).",
    "Calcula el puntaje de cada pregunta usando una escala ponderada de 3 puntos.",
]
for step in proc_steps:
    doc.add_paragraph(step, style="List Number")

doc.add_heading("5.4 Visualización", level=2)
doc.add_paragraph(
    "Una vez cargados los datos, el usuario puede navegar entre las páginas de Resultados, "
    "Comparación y Diagnóstico General para explorar los datos mediante tablas, gráficos "
    "interactivos de Plotly (barras, barras apiladas, circulares, radar, mapas de calor) "
    "y exportar resultados individuales a CSV."
)

# ══════════════════════════════════════════════════════════════════════════
# 6. FORMATO ESPERADO DE ARCHIVOS EXCEL
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("6. Formato Esperado de Archivos Excel", level=1)
doc.add_paragraph(
    "Los archivos Excel deben ser exportados desde la plataforma Moodle en formato .xlsx. "
    "La aplicación espera una estructura muy específica basada en el formato de exportación "
    "estándar de Moodle. A continuación se detalla la estructura requerida."
)

doc.add_heading("6.1 Formato General", level=2)
doc.add_paragraph(
    "El archivo debe contener una única hoja (sheet) con la siguiente estructura:"
)

doc.add_heading("Filas de metadatos (filas 1-13)", level=3)
meta_table = doc.add_table(rows=14, cols=3)
meta_table.style = "Light Shading Accent 1"
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_headers = ["Fila", "Contenido Esperado", "Propósito"]
for i, h in enumerate(meta_headers):
    meta_table.rows[0].cells[i].text = h
meta_data = [
    (
        "1",
        "EVALUACIÓN DE LOS CURSANTES A MÓDULOS DE PROGRAMAS VIRTUALES",
        "Título (se ignora)",
    ),
    ("2", "(vacía)", "Separador"),
    ("3", "CARRERA: <nombre>", "Se extrae el nombre de la carrera"),
    ("4", "SEDE: <nombre>", "Se extrae la sede"),
    ("5", "MODULO: <nombre>", "Se extrae el nombre del módulo"),
    ("6", "GRUPO: <identificador>", "Se extrae el grupo"),
    ("7", "DOCENTE: <nombre>", "Se extrae el nombre del docente/evaluado"),
    ("8", "FECHA INICIO: <fecha>", "Se extrae la fecha de inicio"),
    ("9", "FECHA FIN: <fecha>", "Se extrae la fecha de fin"),
    ("10", "RESPUESTAS ENVIADAS: <número>", "Se extrae la cantidad de respuestas"),
    ("11", "<timestamp>", "Se ignora"),
    ("12", "Preguntas: <número>", "Se ignora"),
    ("13", "(vacía o número)", "Se ignora"),
]
for i, (fila, cont, prop) in enumerate(meta_data, 1):
    meta_table.rows[i].cells[0].text = fila
    meta_table.rows[i].cells[1].text = cont
    meta_table.rows[i].cells[2].text = prop

doc.add_paragraph()

doc.add_heading("Fila de encabezados (fila 14)", level=3)
doc.add_paragraph(
    'La fila 14 debe contener "Etiqueta" en la columna A. Las columnas esperadas son:'
)
hdr_data = [
    ("A", "Etiqueta", "Se ignora (se usa como marcador de inicio)"),
    ("B", "Pregunta", "Texto de la pregunta"),
    ("C", "Excelentes", 'Se mapea a "Sí" (3 puntos)'),
    ("D", "%", "Porcentaje (se ignora)"),
    ("E", "Muy Buenos", 'Se mapea a "A veces" (2 puntos)'),
    ("F", "%", "Porcentaje (se ignora)"),
    ("G", "Bueno", 'Se mapea a "No" (1 punto)'),
    ("H", "Regular", "Siempre None (se ignora)"),
    ("I", "%", "Se ignora"),
    ("J", "Insuficiente", "Siempre None (se ignora)"),
    ("K", "%", "Se ignora"),
]
hdr_table = doc.add_table(rows=1 + len(hdr_data), cols=3)
hdr_table.style = "Light Shading Accent 1"
hdr_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Columna", "Encabezado", "Uso en el sistema"]):
    hdr_table.rows[0].cells[i].text = h
for i, (col, enc, uso) in enumerate(hdr_data, 1):
    hdr_table.rows[i].cells[0].text = col
    hdr_table.rows[i].cells[1].text = enc
    hdr_table.rows[i].cells[2].text = uso

doc.add_paragraph()

doc.add_heading("Filas de preguntas (desde fila 15 en adelante)", level=3)
doc.add_paragraph(
    "Cada pregunta debe estar numerada. El sistema reconoce preguntas que comienzan con "
    'un número seguido de punto o guión (ej: "1. ¿El programa...?" o "1.- ¿El supervisor...?"). '
    "El formato de número es: ^(\\d+)[\\.\\-]"
)
doc.add_paragraph(
    "Los valores numéricos se toman de las columnas C (Excelentes), E (Muy Buenos) y G (Bueno). "
    "Los valores deben ser enteros (conteo de respuestas)."
)

doc.add_heading("Sección de comentarios", level=3)
doc.add_paragraph(
    'Cuando el texto de una pregunta contiene las palabras "observaciones" o "comentarios" '
    "(insensible a mayúsculas/minúsculas), se considera el inicio de la sección de comentarios. "
    "A partir de ese punto, todo el texto en las columnas C en adelante se recolecta como "
    "comentarios, filtrando aquellos que coinciden con palabras vacías como:"
)
stopwords = [
    "ninguno",
    "ninguna",
    "n/a",
    "n/a.",
    "na",
    "no",
    "si",
    "sin observaciones",
    "sin comentarios",
    "sin novedad",
    "no hay observaciones",
    "no hay comentarios",
    "todo bien",
    "todo correcto",
    "bien",
    "ok",
    "ok.",
    "excelente",
    "muy bien",
    "bueno",
    "regular",
    "malo",
]
doc.add_paragraph(", ".join(stopwords))

doc.add_heading("6.2 Requisitos y Restricciones", level=2)
reqs = [
    "El archivo debe ser .xlsx (no .xls).",
    "Debe contener una sola hoja de cálculo. Solo se procesa la hoja activa.",
    "Los metadatos deben usar el formato CLAVE: valor exactamente como se indica (con dos puntos).",
    "Las preguntas deben estar numeradas secuencialmente comenzando en 1.",
    "Los valores de respuesta deben ser números enteros.",
    "La codificación de caracteres debe ser UTF-8 o Latin-1 (Moodle exporta en UTF-8 generalmente).",
]
for r in reqs:
    doc.add_paragraph(r, style="List Bullet")

doc.add_heading("6.3 Cantidad Máxima de Preguntas por Tipo", level=2)
q_table = doc.add_table(rows=8, cols=3)
q_table.style = "Light Shading Accent 1"
q_table.alignment = WD_TABLE_ALIGNMENT.CENTER
q_headers = ["Evaluado", "Evaluador", "Máx. Preguntas"]
for i, h in enumerate(q_headers):
    q_table.rows[0].cells[i].text = h
q_data = [
    ("Docente", "Estudiante", "17"),
    ("Docente", "Equipo de la Carrera", "13"),
    ("Facilitador", "Supervisor", "11"),
    ("Facilitador", "Estudiante", "8"),
    ("Supervisor", "Equipo Pedagógico", "14"),
    ("Supervisor", "Facilitador", "11"),
    ("Asignatura", "Docente", "14"),
]
for i, (ev, eval_, max_q) in enumerate(q_data, 1):
    q_table.rows[i].cells[0].text = ev
    q_table.rows[i].cells[1].text = eval_
    q_table.rows[i].cells[2].text = max_q

doc.add_paragraph()
doc.add_paragraph(
    "Nota: Si un archivo contiene más preguntas que el máximo configurado, las preguntas "
    "excedentes se ignoran silenciosamente."
)

# ══════════════════════════════════════════════════════════════════════════
# 7. MODELO DE PUNTUACIÓN
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("7. Modelo de Puntuación", level=1)

doc.add_heading("7.1 Escala de Valoración", level=2)
doc.add_paragraph("Cada respuesta se valora en una escala de 3 puntos:")
scale_table = doc.add_table(rows=4, cols=3)
scale_table.style = "Light Shading Accent 1"
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(["Respuesta Original (Moodle)", "Categoría Interna", "Valor"]):
    scale_table.rows[0].cells[i].text = h
scale_data = [
    ("Excelentes", "Sí", "3 puntos"),
    ("Muy Buenos", "A veces", "2 puntos"),
    ("Bueno", "No", "1 punto"),
]
for i, (orig, cat, val) in enumerate(scale_data, 1):
    scale_table.rows[i].cells[0].text = orig
    scale_table.rows[i].cells[1].text = cat
    scale_table.rows[i].cells[2].text = val

doc.add_paragraph()

doc.add_heading("7.2 Cálculo del Puntaje por Pregunta", level=2)
doc.add_paragraph(
    "Para cada pregunta, el puntaje se calcula como el promedio ponderado:"
)
p = doc.add_paragraph()
run = p.add_run("Puntaje = (Sí × 3 + A veces × 2 + No × 1) ÷ (Sí + A veces + No)")
run.italic = True

doc.add_paragraph(
    "El resultado se redondea a 2 decimales y se limita a un máximo de 3.0. "
    "Si el total de respuestas es 0, el puntaje se asigna como 0.0."
)

doc.add_heading("7.3 Puntaje por Dimensión", level=2)
doc.add_paragraph(
    "Cada pregunta pertenece a una dimensión pedagógica. El puntaje de una dimensión "
    "es el promedio simple de los puntajes de todas las preguntas que la componen."
)

doc.add_heading("7.4 Puntaje General", level=2)
doc.add_paragraph(
    "El puntaje general de una evaluación es el promedio de todos los puntajes de preguntas "
    "o, equivalentemente, el promedio de los puntajes de todas las dimensiones."
)

doc.add_heading("7.5 Codificación por Colores", level=2)
doc.add_paragraph("Los resultados se visualizan con colores de semáforo:")
color_table = doc.add_table(rows=4, cols=2)
color_table.style = "Light Shading Accent 1"
color_table.alignment = WD_TABLE_ALIGNMENT.CENTER
color_table.rows[0].cells[0].text = "Categoría"
color_table.rows[0].cells[1].text = "Color"
color_table.rows[1].cells[0].text = "Sí"
color_table.rows[1].cells[1].text = "Verde (#1a9641)"
color_table.rows[2].cells[0].text = "A veces"
color_table.rows[2].cells[1].text = "Naranja (#f39c12)"
color_table.rows[3].cells[0].text = "No"
color_table.rows[3].cells[1].text = "Rojo (#d7191c)"

# ══════════════════════════════════════════════════════════════════════════
# 8. DIMENSIONES EVALUADAS POR ROL
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("8. Dimensiones Evaluadas por Rol", level=1)

dims = {
    "Docente": [
        "Planificación e inicio de la sesión",
        "Desarrollo didáctico de los contenidos",
        "Interacción y participación estudiantil",
        "Uso pedagógico de recursos tecnológicos",
        "Evaluación formativa y retroalimentación",
        "Gestión del tiempo y clima de aprendizaje",
    ],
    "Facilitador": [
        "Apoyo y coordinación",
        "Seguimiento académico",
        "Evaluación y retroalimentación",
        "Comunicación efectiva",
        "Actitud profesional y gestión",
    ],
    "Supervisor": [
        "Coordinación y planificación",
        "Seguimiento y supervisión",
        "Comunicación",
        "Gestión de recursos y evaluación",
        "Actitud profesional",
    ],
    "Asignatura": [
        "Inducción y planificación",
        "Comunicación y apoyo",
        "Desarrollo curricular",
        "Recursos y plataforma",
        "Logros y satisfacción",
    ],
}

for role, dimensions in dims.items():
    doc.add_heading(f"{role}", level=2)
    for d in dimensions:
        doc.add_paragraph(d, style="List Bullet")

# ══════════════════════════════════════════════════════════════════════════
# 9. PÁGINAS Y FUNCIONALIDADES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("9. Páginas y Funcionalidades", level=1)

doc.add_heading("9.1 Cargar Datos", level=2)
doc.add_paragraph("Es la página principal al iniciar la aplicación. Permite:")
bullets = [
    "Seleccionar el tipo de evaluado mediante pestañas (Docente, Facilitador, Supervisor, Asignatura).",
    "Dentro de cada tipo, seleccionar el evaluador específico.",
    "Subir archivos .xlsx mediante un cargador de archivos de Streamlit.",
    "Visualizar una tabla resumen con los módulos cargados, incluyendo columnas de: módulo, carrera, docente, respuestas, fecha de inicio, fecha de fin y estado.",
    "Eliminar evaluaciones individuales mediante un botón.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("9.2 Resultados (Individual)", level=2)
doc.add_paragraph("Permite explorar una evaluación en detalle:")
bullets = [
    "Seleccionar tipo de evaluado, evaluador y módulo mediante menús desplegables.",
    "Tarjeta de metadatos con información de la evaluación.",
    "Indicadores KPI: puntaje general, respuestas recibidas, cantidad de preguntas.",
    "Puntajes por dimensión con barras de colores.",
    "Tabla detallada de preguntas con conteos (Sí, A veces, No) y puntaje.",
    "Gráfico de barras apiladas al 100% mostrando la distribución de respuestas por pregunta.",
    "Gráfico de barras de puntajes por pregunta con línea de referencia.",
    "Gráfico circular (pie chart) de la distribución general de respuestas.",
    "Gráfico de radar con puntajes por dimensión.",
    "Sección de comentarios de texto libre.",
    "Botón para descargar resultados en CSV.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("9.3 Comparación", level=2)
doc.add_paragraph("Agrupa y compara evaluaciones por módulo:")
bullets = [
    "Seleccionar tipo de evaluado y evaluador.",
    "Tabla de comparación con puntajes promediados por módulo para cada dimensión.",
    "Ranking de módulos ordenado por puntaje general descendente.",
    "Promedios por tipo de evaluador.",
    "Gráfico de barras de puntajes por dimensión y tipo.",
    "Gráfico de radar comparativo entre módulos.",
    "Tabla detallada con el promedio de cada pregunta por módulo.",
    "Mapa de calor (heatmap) de puntajes por módulo y dimensión.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

doc.add_heading("9.4 Diagnóstico General", level=2)
doc.add_paragraph("Panel consolidado con todas las evaluaciones cargadas:")
bullets = [
    "Métrica global: promedio general de todas las evaluaciones.",
    "Total de evaluaciones cargadas.",
    "Total de respuestas recibidas (suma de todas las evaluaciones).",
    "Ranking completo de todas las evaluaciones ordenadas por puntaje.",
    "Análisis de puntajes por dimensión desglosados por tipo de evaluado.",
    "Tabla de fortalezas (puntaje ≥ 2.8) con resaltado en verde.",
    "Tabla de debilidades (puntaje ≤ 2.0) con resaltado en rojo.",
    "Estadísticas descriptivas: mediana, desviación estándar, mínimo, máximo por dimensión.",
]
for b in bullets:
    doc.add_paragraph(b, style="List Bullet")

# ══════════════════════════════════════════════════════════════════════════
# 10. LIMITACIONES Y CONSIDERACIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("10. Limitaciones y Consideraciones", level=1)

doc.add_heading("10.1 Limitaciones Funcionales", level=2)
limitations = [
    "Escala reducida: Aunque Moodle exporta 5 niveles (Excelentes, Muy Buenos, Bueno, Regular, Insuficiente), "
    "el sistema solo utiliza los 3 primeros. Las columnas Regular e Insuficiente siempre se ignoran porque "
    "los archivos de muestra no contienen valores en esas columnas.",
    "Datos en memoria: Toda la información se almacena en st.session_state. Al recargar la página o "
    "cerrar el navegador, todos los datos cargados se pierden. No hay persistencia en base de datos ni "
    "en disco.",
    "Sin exportación masiva: Solo la página de Resultados permite descargar CSV. Las páginas de "
    "Comparación y Diagnóstico General no ofrecen exportación de datos.",
    "Parseo frágil de metadatos: La extracción usa str.startswith() que es sensible a espacios en blanco "
    'adicionales. Un espacio extra después de los dos puntos ("CARRERA: " vs "CARRERA:") podría romper '
    "la extracción.",
    "Límite de preguntas: Si un archivo contiene más preguntas que el máximo configurado para ese tipo "
    "de evaluado, las preguntas excedentes se descartan silenciosamente sin notificación al usuario.",
    "Formato de pregunta rígido: Las preguntas deben comenzar con un número seguido de punto o guión. "
    "Preguntas sin este formato reciben un número de fallback incorrecto.",
    "Filtro de comentarios limitado: La lista de stop-words para filtrar comentarios es fija, "
    "sensible a mayúsculas/minúsculas (solo minúsculas), y puede dejar pasar comentarios irrelevantes "
    "o filtrar comentarios válidos.",
    "Una sola hoja: Solo se procesa la hoja activa del archivo Excel. Archivos con múltiples hojas "
    "no se procesan completamente.",
    "Sin validación de entrada: No hay verificación de columnas faltantes, formato incorrecto, "
    "archivos vacíos o archivos corruptos. La función safe() oculta silenciosamente errores de "
    "conversión de tipos.",
    "Configuración hard-codeada: Todos los tipos de evaluación, dimensiones, colores y textos de "
    "preguntas están hard-codeados en el código fuente. Agregar un nuevo tipo de encuesta requiere "
    "modificar el código.",
    "Sin soporte multi-idioma: Toda la interfaz está en español sin soporte para internacionalización.",
]
for i, lim in enumerate(limitations, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(lim)

doc.add_heading("10.2 Consideraciones Técnicas", level=2)
tech_cons = [
    "El archivo de la aplicación se llama app_completa.py, pero el README.md indica ejecutar "
    "streamlit run app.py. El comando correcto es streamlit run app_completa.py.",
    "La función de deduplicación (dedup) solo está activa para la combinación Docente → Estudiante. "
    "Esto sugiere que en algunos casos Moodle exporta preguntas duplicadas para ese tipo de evaluación.",
    "Los gráficos de radar tienen escalas fijas [0, 3] y etiquetas hard-codeadas.",
    "No se requiere conexión a internet para usar la aplicación una vez cargada en el navegador.",
    "El proyecto no utiliza variables de entorno ni archivos .env.",
]
for c in tech_cons:
    doc.add_paragraph(c, style="List Bullet")

# ══════════════════════════════════════════════════════════════════════════
# 11. RECOMENDACIONES
# ══════════════════════════════════════════════════════════════════════════
doc.add_heading("11. Recomendaciones", level=1)
recs = [
    "Validar los archivos Excel antes de procesarlos: verificar que tengan la estructura esperada, "
    "que los encabezados coincidan y que los datos no estén corruptos.",
    "Agregar persistencia de datos (base de datos SQLite o archivos JSON) para evitar la pérdida "
    "de información al recargar la página.",
    "Implementar exportación de datos en las páginas de Comparación y Diagnóstico General.",
    "Ampliar la escala de valoración para utilizar los 5 niveles de Moodle si los datos están "
    "disponibles.",
    "Mejorar el parseo de metadatos usando expresiones regulares más robustas en lugar de startswith().",
    "Agregar notificaciones al usuario cuando se exceda el límite de preguntas.",
    "Externalizar la configuración (tipos de evaluación, dimensiones, colores) a un archivo JSON "
    "o YAML para facilitar la personalización sin modificar código.",
    "Implementar logging para facilitar la depuración de errores al procesar archivos.",
    "Agregar pruebas unitarias para las funciones de parseo y cálculo de puntuaciones.",
    "Corregir el README.md para que indique el nombre correcto del archivo principal.",
]
for i, rec in enumerate(recs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"{i}. ")
    run.bold = True
    p.add_run(rec)

# ── Guardar ──────────────────────────────────────────────────────────────
output_path = "Documentacion_App_Encuestas.docx"
doc.save(output_path)
print(f"Documento generado exitosamente: {output_path}")
